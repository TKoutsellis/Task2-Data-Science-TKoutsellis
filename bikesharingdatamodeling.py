#bikesharingdatamodeling.py
# -*- coding: utf-8 -*-
"""
Created on Sat Nov 30 17:02:12 2019

@author: Koutsellis Themistoklis

Description
-----------
Creates all plots for task2
Saves them to Report.docx

"""

import io
import sys
import pandas as pd
import numpy as np
import seaborn as sns
import matplotlib.pyplot as plt
from docx import Document # pip install python-docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sklearn.model_selection import train_test_split
from sklearn import model_selection
from sklearn.tree import DecisionTreeRegressor
from sklearn.linear_model import LinearRegression, \
                                 Ridge, \
                                 Lasso
from sklearn.ensemble import BaggingRegressor, \
                             GradientBoostingRegressor,\
                             RandomForestRegressor

sns.set(style="darkgrid")

# initializing the Report.docx file
DOCUMENT = Document()
DOCUMENT.add_heading('Report', 0)


def plot_decorator(func, *args, **kwargs):
    '''
    decorator applied to all plot function.
    handles the exception.
    '''
    def wrapper(*args, **kwargs):
        try:
            plt.figure(figsize=(16, 9))
            func(*args, **kwargs)
            memfile = io.BytesIO()
            plt.savefig(memfile)
            # Don't show plot in (Spyder)Console. Just save it in Report.docx
            plt.close()
            # save plot to Report.docx
            global DOCUMENT
            DOCUMENT.add_picture(memfile, width=Inches(6.0)) #width=Inches(6.3)
            paragraph = DOCUMENT.add_paragraph('')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run('Figure .')
            run.bold = True
            paragraph.add_run([' method: ', func.__name__, '()'])
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            DOCUMENT.save('Report.docx')
            memfile.close()
        except Exception as ex:
            print(func)
            print("\n* Operation failed with the following error: ", ex, '\n')
            sys.exit(1)
    return wrapper


class BikeSharingDataModeling:
    '''
    TestCase class to test fetchimages.fetch_images() function.
    '''

    def __init__(self, csv_file_path, *args, **kwargs):
        try:
            # Making a list of extra non-standard missing value types
            missing_values = ["n/a", "na", "--", "-", " ", "."]
            self.csv_file = pd.read_csv(csv_file_path, na_values=missing_values)
        except Exception as ex:
            print("\n* Operation failed with the following error: ", ex, '\n')
            sys.exit(1)

    def check_for_missing_data(self, *args, **kwargs):
        '''
        checks for missing data in the csv file
        '''
        pd.options.mode.use_inf_as_na = True
        if not self.csv_file.isnull().values.any():
            print('No missing data')
        else:
            print('!There are {} missing data!'.format(self.csv_file.isnull().sum().sum()))
            sys.exit(1)

    def brief_overview_of_the_data(self, *args, **kwargs):
        '''
        Displays the 10 first lines of the DataFrame
        '''
        print()
        print('Data overview:')
        print(self.csv_file.head())

    def drop_columns_from_data(self, drop_list, *args, **kwargs):
        '''
        Drops out the fields specified in drop_list
        '''
        try:
            self.csv_file = self.csv_file.drop(drop_list, axis=1)
        except Exception as ex:
            print("\n* Operation failed with the following error: ", ex, '\n')
            sys.exit(1)

    # Figure 1
    @plot_decorator
    def cnt_vs_season(self, *args, **kwargs):
        '''
        plot: season vs (casual, regular, cnt) variables
        '''
        data = self.csv_file.loc[:, ['season', 'cnt', 'casual', 'registered']]
        plot = sns.lineplot(x='season',
                            y='value',
                            hue='variable',
                            data=pd.melt(data, ['season']),
                            markers=["o", ".", "*"],
                            )
        plt.setp(plot.get_legend().get_texts(), fontsize='16')
        plt.setp(plot.get_legend().get_title(), fontsize='18')
        plt.title("Count of rental bikes vs Seasons ", fontsize=20)
        plt.xlabel("Season [1:spring, 2:summer, 3:fall, 4:winter] ", fontsize=18)
        plt.ylabel("Rate: Counts/hours", fontsize=18)

    # Figure 2
    @plot_decorator
    def cnt_vs_month(self, *args, **kwargs):
        '''
        plot: month vs (casual, regular, cnt) variables
        '''

        data = self.csv_file.loc[:, ['mnth', 'cnt', 'casual', 'registered']]
        plot = sns.lineplot(x='mnth',
                            y='value',
                            hue='variable',
                            data=pd.melt(data, ['mnth']),
                            legend="brief",
                            markers=["o", ".", "*"],
                            )
        plt.setp(plot.get_legend().get_texts(), fontsize='16')
        plt.setp(plot.get_legend().get_title(), fontsize='18')
        plt.title("Count of rental bikes vs Months", fontsize=20)
        plt.xlabel("Month ", fontsize=18)
        plt.ylabel("Rate: Counts/hours", fontsize=18)

    # Figure 3
    @plot_decorator
    def cnt_vs_hour(self, *args, **kwargs):
        '''
        plot: hour vs (casual, regular, cnt) variables
        '''
        data = self.csv_file.loc[:, ['hr', 'cnt', 'casual', 'registered']]
        plot = sns.lineplot(x='hr',
                            y='value',
                            hue='variable',
                            data=pd.melt(data, ['hr']),
                            legend="brief",
                            markers=["o", ".", "*"],
                            )
        plt.setp(plot.get_legend().get_texts(), fontsize='16')
        plt.setp(plot.get_legend().get_title(), fontsize='18')
        plt.title("Count of rental bikes vs Hours", fontsize=20)
        plt.xlabel("Hour", fontsize=18)
        plt.ylabel("Rate: Counts/hours", fontsize=18)

    # Figure 4
    @plot_decorator
    def cnt_vs_holiday(self, *args, **kwargs):
        '''
        plot: cnt vs weathersit
        '''
        data = self.csv_file.loc[:, ['holiday', 'cnt', 'casual', 'registered']]
        data = pd.melt(data, ['holiday'])
        sns.catplot(x="holiday",
                    y="value",
                    hue="variable",
                    legend="brief",
                    kind="point",
                    data=data,
                    )
        #plot.despine(left=True)
        plt.title("Count of rental bikes vs weathersit", fontsize=20)
        plt.xlabel("Weathersit: [1: Clear, 2: Mist, 3: Light Snow/Rain, 4: Heavy Rain/Snow]",
                   fontsize=18)
        plt.ylabel("Rate: Counts/hours", fontsize=18)

    # Figure 5
    @plot_decorator
    def cnt_vs_weekday(self, *args, **kwargs):
        '''
        plot: weekday vs (casual, regular, cnt) variables
        '''
        data = self.csv_file.loc[:, ['weekday', 'cnt', 'casual', 'registered']]
        plot = sns.lineplot(x='weekday',
                            y='value',
                            hue='variable',
                            data=pd.melt(data, ['weekday']),
                            legend="brief",
                            markers=["o", ".", "*"],
                            )
        plt.setp(plot.get_legend().get_texts(), fontsize='16')
        plt.setp(plot.get_legend().get_title(), fontsize='18')
        plt.title("Count of rental bikes vs holiday", fontsize=20)
        plt.xlabel("weekday:", fontsize=18)
        plt.ylabel("Rate: Counts/hours", fontsize=18)

    # Figure 6
    @plot_decorator
    def cnt_vs_workingday(self, *args, **kwargs):
        '''
        plot: cnt vs weathersit
        '''
        data = self.csv_file.loc[:, ['workingday', 'cnt', 'casual', 'registered']]
        data = pd.melt(data, ['workingday'])
        sns.catplot(x="workingday",
                    y="value",
                    hue="variable",
                    legend="brief",
                    kind="point",
                    data=data,
                    )
        plt.title("Count of rental bikes vs workingday", fontsize=20)
        plt.xlabel("workingday:", fontsize=18)
        plt.ylabel("Rate: Counts/hours", fontsize=18)

    # Figure 7
    @plot_decorator
    def cnt_vs_hours_per_month(self, *args, **kwargs):
        '''
        plot: cnt vs hour (per month)
        '''
        data = self.csv_file.loc[:, ['hr', 'mnth', 'cnt']]
        plot = sns.lineplot(x='hr',
                            y='cnt',
                            hue='mnth',
                            data=data,
                            legend="brief",
                            markers=["o", ".", "*"],
                            )
        plt.setp(plot.get_legend().get_texts(), fontsize='16')
        plt.setp(plot.get_legend().get_title(), fontsize='18')
        plt.title("Count of rental bikes vs hour per month", fontsize=20)
        plt.xlabel("hour", fontsize=18)
        plt.ylabel("Rate: Counts/hours", fontsize=18)

    # Figure 8
    @plot_decorator
    def cnt_vs_hours_per_weekday(self, *args, **kwargs):
        '''
        plot: cnt vs hour (per weekday)
        '''
        data = self.csv_file.loc[:, ['hr', 'weekday', 'cnt']]
        plot = sns.lineplot(x='hr',
                            y='cnt',
                            hue='weekday',
                            data=data,
                            legend="brief",
                            markers=["o", ".", "*"],
                            )
        plt.setp(plot.get_legend().get_texts(), fontsize='16')
        plt.setp(plot.get_legend().get_title(), fontsize='18')
        plt.title("Count of rental bikes  vs hour per weekday", fontsize=20)
        plt.xlabel("weekday", fontsize=18)
        plt.ylabel("Rate: Counts/hours", fontsize=18)

    # Figure 9
    @plot_decorator
    def cnt_vs_hours_per_working_day(self, *args, **kwargs):
        '''
        plot: cnt vs hour (per weekday)
        '''
        data = self.csv_file.loc[:, ['hr', 'workingday', 'cnt']]
        plot = sns.lineplot(x='hr',
                            y='cnt',
                            hue='workingday',
                            data=data,
                            legend="brief",
                            markers=["o", ".", "*"],
                            )
        plt.setp(plot.get_legend().get_texts(), fontsize='16')
        plt.setp(plot.get_legend().get_title(), fontsize='18')
        plt.title("Count of rental bikes[cnt] vs hour per workingday", fontsize=20)
        plt.xlabel("hour", fontsize=18)
        plt.ylabel("Rate: Counts/hours", fontsize=18)

    # Figure 10
    @plot_decorator
    def corr_matrix(self, *args, **kwargs):
        '''
        plot: corr. matrix
        '''
        correlations = self.csv_file.corr()
        mask = np.array(correlations)
        mask[np.tril_indices_from(mask)] = False
        sns.heatmap(correlations,
                    mask=mask,
                    vmax=1,
                    vmin=-1,
                    square=True,
                    annot=True,
                    )
        plt.title("Correlation Matrix", fontsize=20)

    # Figure 11
    @plot_decorator
    def cnt_violin_weathersit(self, *args, **kwargs):
        '''
        plot: cnt vs weathersit
        '''
        data = self.csv_file.loc[:, ['weathersit', 'cnt']]
        # Show each distribution with both violins and points
        plot = sns.violinplot(data=data,
                              inner="points",
                              x='weathersit',
                              y='cnt',
                              hue='weathersit',
                              )
        plt.setp(plot.get_legend().get_texts(), fontsize='16')
        plt.setp(plot.get_legend().get_title(), fontsize='18')
        plt.title("Count of rental bikes[cnt] vs weathersit", fontsize=20)
        plt.xlabel("Weathersit: [1: Clear, 2: Mist, 3: Light Snow/Rain, 4: Heavy Rain/Snow]",
                   fontsize=18)
        plt.ylabel("Rate: Counts/hours", fontsize=18)

    # Figure 12
    @plot_decorator
    def cnt_vs_hum(self, *args, **kwargs):
        '''
        plot: cnt vs humidity
        '''
        data = self.csv_file.loc[:, ['hum', 'cnt']]
        sns.lineplot(x='hum',
                     y='cnt',
                     data=data,
                     markers=["*"],
                    )
        plt.title("Count of rental bikes[cnt] vs humidity", fontsize=20)
        plt.xlabel("humidity", fontsize=18)
        plt.ylabel("Rate: Counts/hours", fontsize=18)

    # Figure 13
    @plot_decorator
    def cnt_vs_temp(self, *args, **kwargs):
        '''
        plot: cnt vs temperature
        '''
        data = self.csv_file.loc[:, ['temp', 'cnt']]
        sns.lineplot(x='temp',
                     y='cnt',
                     data=data,
                     markers=["*"],
                     )
        plt.title("Count of rental bikes[cnt] vs temperature", fontsize=20)
        plt.xlabel("temperature", fontsize=18)
        plt.ylabel("Rate: Counts/hours", fontsize=18)

    # Figure 14
    @plot_decorator
    def cnt_vs_windspeed(self, *args, **kwargs):
        '''
        plot: cnt vs windspeed
        '''
        data = self.csv_file.loc[:, ['windspeed', 'cnt']]
        sns.lineplot(x='windspeed',
                     y='cnt',
                     data=data,
                     markers=["*"],
                     )
        plt.title("Count of rental bikes[cnt] vs windspeed", fontsize=20)
        plt.xlabel("windspeed", fontsize=18)
        plt.ylabel("Rate: Counts/hours", fontsize=18)

    # Figure 15
    @plot_decorator
    def box_plot_cnt_vs_hours(self, *args, **kwargs):
        '''
        plot: corr. matrix
        '''
        sns.boxplot(x=self.csv_file['hr'],
                    y=self.csv_file['cnt'],
                    )
        plt.title("Boxplot for Count[cnt] Variable", fontsize=20)
        plt.xlabel("hour", fontsize=18)
        plt.ylabel("Count", fontsize=18)

    # Figure 16
    @plot_decorator
    def histagram_cnt_vs_hours(self, *args, **kwargs):
        '''
        plot: corr. matrix
        '''
        sns.distplot(self.csv_file['cnt'])
        plt.title("Counts[cnt] histogram", fontsize=20)
        plt.xlabel("hour", fontsize=18)
        plt.ylabel("Count", fontsize=18)

    def model_selection(self, models, *args, **kwargs):
        '''
        Selects the most effective model from a list of models

        Arguments:
            models -- list of model functions from sklearn
        '''


#        Data dummification:
#        -------------------
#        one-hot encoding: categorical variables
#        are converted into a form that could be
#        provided to ML algorithms to do a better
#        job in prediction.

        dummy_data = self.csv_file

        def data_dummification(df2, column):

            df2 = pd.concat([df2,
                             pd.get_dummies(df2[column],
                             prefix=column,
                             drop_first=True)
                            ],
                            axis=1
                            )
            df2 = df2.drop([column], axis=1)
            return df2

        dummify_cols = ['mnth',
                        'hr',
                        'holiday',
                        'weekday',
                        'workingday',
                        'weathersit',
                        ]
        for col in dummify_cols:
            dummy_data = data_dummification(dummy_data, col)

        # Split the dummy_data into train and test data
        # to train and test the model respectively.
        Y = dummy_data['cnt']
        X = dummy_data.drop(['cnt'], axis=1)
        X_train, X_test, Y_train, Y_test =  \
        train_test_split(X, Y,
                         train_size=0.7,
                         test_size=0.3,
                         )

        ###############################
        #         Chose model         #
        ###############################

#        # Compare regrassions algirithms
#        models = [Ridge(),
#                  Lasso(),
#                  LinearRegression(),
#                  BaggingRegressor(),
#                  DecisionTreeRegressor(),
#                  RandomForestRegressor(),
#                  GradientBoostingRegressor(),
#                  ]

        # initializing list with validation score per model
        cross_val_score_list = []

        def algorithm_test(model, cross_val_score_list):
            '''
            test given model
            Arguments:
                model -- the model
                cross_val_score_list -- list with validation score per model
            Returns:
                cross_val_score_list -- list with validation score per model

            '''
            kfold = model_selection.KFold(n_splits=10)
            predicted = model_selection.cross_val_score(model,
                                                        X_train,
                                                        Y_train,
                                                        cv=kfold,
                                                        scoring='neg_mean_squared_error',
                                                        )
            cross_val_score_list.append(predicted.mean())
            return cross_val_score_list

        # Next snippet needs time. Warn user about that
        print('Selecting the best algorithm')
        print('Please wait ...')
        # Append the mean cross val score to each model
        for model in models:
            cross_val_score_list = algorithm_test(model,
                                                  cross_val_score_list
                                                  )

        # select the model with the smallest abs cross val score
        selected_model_idx = cross_val_score_list.index( \
               max(cross_val_score_list))

        # selected model
        selected_model = models[selected_model_idx]

        # execute RandomForestRegressor
        rf = selected_model
        # Create Y traind data - prediction
        rf.fit(X_train, Y_train)
        Y_pred = rf.predict(X_test)

        # calculate MAD (Mean Absolute Deviation)
        mad = X_train.mad()

        # log print
        print('\n>>>>>> Selected model:\n{}'.format(selected_model))

        print('\nMean Absolute Deviations:')
        print(mad)

if __name__ == '__main__':
    bike_sharing_data = BikeSharingDataModeling('hour.csv')
    drop_fields_list = ['instant',
                        'dteday',
                        'yr',
                        'season',
                        'casual',
                        'registered',
                        'atemp',
                        'windspeed',
                        ]

    bike_sharing_data = BikeSharingDataModeling('hour.csv')
    bike_sharing_data.check_for_missing_data()
    bike_sharing_data.brief_overview_of_the_data()


    # plots
    print('\nCreating all necessary plots')
    print('Plase wait...\n')
    bike_sharing_data.cnt_vs_season()
    bike_sharing_data.cnt_vs_month()
    bike_sharing_data.cnt_vs_hour()
    bike_sharing_data.cnt_vs_holiday()
    bike_sharing_data.cnt_vs_weekday()
    bike_sharing_data.cnt_vs_workingday()
    bike_sharing_data.cnt_vs_hours_per_month()
    bike_sharing_data.cnt_vs_hours_per_weekday()
    bike_sharing_data.cnt_vs_hours_per_working_day()
    bike_sharing_data.corr_matrix()
    bike_sharing_data.cnt_violin_weathersit()
    bike_sharing_data.cnt_vs_hum()
    bike_sharing_data.cnt_vs_temp()
    bike_sharing_data.cnt_vs_windspeed()
    # distribution plots
    bike_sharing_data.box_plot_cnt_vs_hours()
    bike_sharing_data.histagram_cnt_vs_hours()
    # modelling
    bike_sharing_data.drop_columns_from_data(drop_fields_list)
    # List of model to compare
    models = [Ridge(),
              Lasso(),
              LinearRegression(),
              BaggingRegressor(),
              DecisionTreeRegressor(),
              RandomForestRegressor(),
              GradientBoostingRegressor(),
              ]
    # Compare model and pick the most effective one
    bike_sharing_data.model_selection(models)
